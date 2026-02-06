import os
import fitz
import pandas as pd
from docx import Document
from sentinel.tools.smart_index import index_file

try:
    import tabulate
except ImportError:
    tabulate = None

def read_file(path):
    """
    Reads a file safely. Supports PDF, DOCX, XLSX, CSV, TXT, MD, PY.
    """
    if not os.path.exists(path):
        return f"❌ Error: File not found at {path}"

    ext = os.path.splitext(path)[1].lower()

    try:
        # --- PDF Handling ---
        if ext == '.pdf':
            text = ""
            with fitz.open(path) as doc:
                for i, page in enumerate(doc):
                    if i >= 20: break
                    text += page.get_text() + "\n"
            return text if text.strip() else "⚠️ PDF contains no selectable text (It might be an image scan)."

        # --- Word Docs ---
        elif ext == '.docx':
            doc = Document(path)
            return "\n".join([p.text for p in doc.paragraphs])

        # --- Excel / CSV ---
        elif ext in ['.xlsx', '.xls', '.csv']:
            if not tabulate:
                return "❌ Error: 'tabulate' library missing. Run `pip install tabulate`."

            if ext == '.csv':
                df = pd.read_csv(path, nrows=50)
            else:
                df = pd.read_excel(path, nrows=50)

            return df.to_markdown(index=False)

        # --- Code / Plain Text ---
        else:
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(15000)  # Increased limit slightly
                if len(content) == 15000:
                    content += "\n... [Content Truncated]"
                return content

    except Exception as e:
        return f"❌ Error reading file: {e}"

def write_file(path, content):
    """
    Writes content to a file.
    Auto-detects .docx to create valid Word documents.
    """
    try:
        folder = os.path.dirname(path)
        if folder and not os.path.exists(folder):
            os.makedirs(folder)

        ext = os.path.splitext(path)[1].lower()

        if ext == '.docx':
            doc = Document()
            for line in content.split('\n'):
                doc.add_paragraph(line)
            doc.save(path)

        else:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(str(content))

        try:
            index_file(path)
            return f"✅ Saved and Indexed: {path}"
        except Exception as e:
            return f"✅ Saved to {path}, but indexing failed: {e}"

    except Exception as e:
        return f"❌ Write Failed: {e}"