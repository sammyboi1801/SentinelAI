import os
import platform
import subprocess
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _convert_to_pdf(docx_path):
    """
    Converts a .docx file to .pdf using LibreOffice (if available) or basic fallback.
    Note: Perfect .docx -> .pdf conversion usually requires Word installed or LibreOffice headless.
    For this MVP, we will try a pure Python fallback if system tools aren't found.
    """
    try:
        cmd = f"soffice --headless --convert-to pdf \"{docx_path}\" --outdir \"{os.path.dirname(docx_path)}\""
        if os.system(cmd) == 0:
            return docx_path.replace(".docx", ".pdf")

        try:
            from docx2pdf import convert
            convert(docx_path)
            return docx_path.replace(".docx", ".pdf")
        except:
            pass

        return "PDF conversion requires LibreOffice or MS Word installed."
    except Exception as e:
        return f"PDF Error: {e}"


def create_document(filename, blocks):
    """
    Universal Document Generator.

    Args:
        filename (str): Output name (e.g. "Project_Plan")
        blocks (list): A list of dictionaries defining the structure.

    Block Types:
        - {"type": "heading", "text": "My Title", "level": 1}
        - {"type": "paragraph", "text": "Body content...", "bold": False}
        - {"type": "list", "items": ["Bullet 1", "Bullet 2"], "ordered": False}
        - {"type": "table", "rows": [["Header 1", "Header 2"], ["Cell A", "Cell B"]]}
        - {"type": "page_break"}
    """
    try:
        doc = Document()

        # Set default style (Optional)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        for block in blocks:
            b_type = block.get("type", "paragraph")

            # --- HEADING ---
            if b_type == "heading":
                text = block.get("text", "")
                level = block.get("level", 1)
                doc.add_heading(text, level)

            # --- PARAGRAPH ---
            elif b_type == "paragraph":
                text = block.get("text", "")
                p = doc.add_paragraph()
                run = p.add_run(text)
                if block.get("bold"): run.bold = True
                if block.get("italic"): run.italic = True

                align = block.get("align", "left").lower()
                if align == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif align == "justify":
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # --- LISTS ---
            elif b_type == "list":
                items = block.get("items", [])
                style = 'List Number' if block.get("ordered") else 'List Bullet'
                for item in items:
                    doc.add_paragraph(item, style=style)

            # --- TABLES ---
            elif b_type == "table":
                rows = block.get("rows", [])
                if rows:
                    cols = len(rows[0])
                    table = doc.add_table(rows=0, cols=cols)
                    table.style = 'Table Grid'

                    for row_data in rows:
                        row_cells = table.add_row().cells
                        for i, cell_text in enumerate(row_data):
                            if i < len(row_cells):
                                row_cells[i].text = str(cell_text)

            elif b_type == "page_break":
                doc.add_page_break()

        # Save Logic
        if not os.path.exists("drafts"):
            os.makedirs("drafts")

        if not filename.endswith(".docx"): filename += ".docx"
        docx_path = os.path.join("drafts", filename)

        doc.save(docx_path)

        pdf_result = _convert_to_pdf(docx_path)

        return f"Document created: {docx_path} (PDF Status: {pdf_result})"

    except Exception as e:
        return f"Document Factory Error: {e}"